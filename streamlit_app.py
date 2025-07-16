import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import string
from docx.shared import Pt

st.title("CSV to Word Table Generator")

# Configuration options
term = st.selectbox("Select Term", ["SPRING", "MONSOON"])
year = st.text_input("Enter Year", value="2025")
semester = st.text_input("Enter Semester", value="4")
section = st.selectbox("Select Section", list(string.ascii_uppercase))

st.write("Upload one or more CSV files. Each will be added as a table in the generated Word document. If you upload a file named 'cross_table.csv', it will always appear last in the document.")

uploaded_files = st.file_uploader("Upload CSV files", type=["csv"], accept_multiple_files=True)

# Mapping for custom headings
custom_headings = {
    'q0.csv': 'Question 1',
    'q1.csv': 'Question 2',
    'q2.csv': 'Question 3',
    'q3.csv': 'Question 4',
    'q4.csv': 'Question 5',
    'q5.csv': 'Question 6',
    'q6.csv': 'Question 7',
    'q7.csv': 'Question 8',
    'cross_table.csv': 'Deviations across questions',
}

def add_title_paragraph(doc, text):
    p = doc.add_paragraph(text, style='Title')
    for run in p.runs:
        run.font.size = Pt(18)

def add_df_to_doc(doc, df, title):
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        if col == 'Unnamed: 0':
            hdr_cells[i].text = ''
        else:
            hdr_cells[i].text = str(col)
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

if uploaded_files:
    if st.button("Generate Word Document"):
        doc = Document()
        # Add the custom header with smaller title font
        add_title_paragraph(doc, "Department of Computer Science & Engineering")
        add_title_paragraph(doc, "Birla Institute of Technology, Mesra")
        add_title_paragraph(doc, "Faculty Feedback Action Taken Report")
        add_title_paragraph(doc, f"{term.capitalize()} {year}")
        doc.add_paragraph("")
        doc.add_paragraph(f"Program: Computer Science & Engineering (Semester {semester}: Section {section})", style='Heading 1')
        doc.add_paragraph("Question Wise Feedback", style='Heading 2')
        doc.add_paragraph("")
        # Add the term/semester/section line before tables
        doc.add_paragraph(f"{term.capitalize()} {year}, Semester {semester}: Section {section}", style='Heading 2')
        doc.add_paragraph("")
        # Separate cross_table.csv from others
        cross_table_file = None
        other_files = []
        for uploaded_file in uploaded_files:
            if uploaded_file.name.lower() == 'cross_table.csv':
                cross_table_file = uploaded_file
            else:
                other_files.append(uploaded_file)
        # Add other files first
        for uploaded_file in other_files:
            df = pd.read_csv(uploaded_file)
            heading = custom_headings.get(uploaded_file.name.lower(), uploaded_file.name)
            add_df_to_doc(doc, df, heading)
            doc.add_paragraph("")
        # Add cross_table.csv last if present
        if cross_table_file:
            df = pd.read_csv(cross_table_file)
            heading = custom_headings.get(cross_table_file.name.lower(), cross_table_file.name)
            add_df_to_doc(doc, df, heading)
            doc.add_paragraph("")
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.success("Word document generated!")
        st.download_button(
            label="Download Word Document",
            data=buffer,
            file_name="CS.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ) 